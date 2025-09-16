"""
Parse attachments and normalize them into pricing rows (or text blocks for the LLM).

Supports:
- Excel/CSV → read with pandas/openpyxl and map common columns to a standard schema
- PDF/DOCX → extract text; return as "texts" for the LLM to interpret

Output contract:
- For spreadsheet-like inputs: normalized rows (dicts) ready for diff/snapshot.
- For PDF/DOCX: raw text blocks (strings) that llm/extractor.py can process.

Tip: If a new supplier has unusual column names, extend the column mapping here
instead of hacking the rest of the pipeline.
"""


import re
from io import BytesIO
from typing import List, Dict, Any

import pandas as pd

try:
    import docx  # python-docx
except ImportError:
    docx = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None


# --------- Hjälpfunktioner ---------
def _norm_col(s: str) -> str:
    return re.sub(r"[^a-z0-9]", "", s.strip().lower())


def _to_float(val):
    if val is None:
        return None
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val)
    s = s.replace("€", "").replace("$", "").replace("£", "")
    s = s.replace("\u00a0", " ").strip()
    s = s.replace(",", ".")
    s = re.sub(r"[A-Za-z]", "", s)
    s = re.sub(r"[^0-9.\-]", "", s)
    if s in ("", ".", "-"):
        return None
    try:
        return float(s)
    except Exception:
        return None


# synonymer -> fältnamn i vårt schema
_COL_MAP = {
    "country": "country",
    "countryiso": "country_iso",
    "countrycode": "country_code",
    "cc": "country_code",
    "operator": "operator",
    "network": "network",
    "mcc": "mcc",
    "mnc": "mnc",
    "imsi": "imsi",
    "nnc": "nnc",
    "numbertype": "number_type",
    "destination": "destination",
    "previousrate": "previous_rate",
    "oldrate": "old_price",
    "currentrate": "current_rate",
    "newprice": "new_price",
    "rate": "price",
    "rateeur": "price",
    "price": "price",
    "currency": "currency",
    "variation": "variation",
    "change": "variation",
    "valid": "effective_from",
    "effectivedate": "effective_from",
    "effectivefrom": "effective_from",
    "count": "count",
    "cost": "cost",
    "costeur": "cost",
    "productcategory": "product_category",
    "mccmnc": "destination",  # ibland används som sammanslaget
}


def _map_columns(df: pd.DataFrame) -> pd.DataFrame:
    rename = {}
    for c in df.columns:
        key = _norm_col(str(c))
        tgt = _COL_MAP.get(key)
        if tgt:
            rename[c] = tgt
    if rename:
        df = df.rename(columns=rename)
    return df


def _row_to_schema(row: dict, default_currency=None) -> dict:
    out = {
        "provider": None,
        "country": row.get("country"),
        "country_iso": row.get("country_iso"),
        "country_code": row.get("country_code"),
        "operator": row.get("operator"),
        "network": row.get("network"),
        "mcc": (str(row.get("mcc")) if row.get("mcc") is not None else None),
        "mnc": (str(row.get("mnc")) if row.get("mnc") is not None else None),
        "imsi": (str(row.get("imsi")) if row.get("imsi") is not None else None),
        "nnc": (str(row.get("nnc")) if row.get("nnc") is not None else None),
        "number_type": row.get("number_type"),
        "destination": (str(row.get("destination")) if row.get("destination") is not None else None),
        "previous_rate": _to_float(row.get("previous_rate")) if row.get("previous_rate") is not None else _to_float(row.get("old_price")),
        "old_price": _to_float(row.get("old_price")),
        "current_rate": _to_float(row.get("current_rate")),
        "new_price": _to_float(row.get("new_price")),
        "price": _to_float(row.get("price")),
        "currency": row.get("currency") or default_currency,
        "variation": (str(row.get("variation")).strip().lower() if row.get("variation") else None),
        "effective_from": (str(row.get("effective_from")) if row.get("effective_from") is not None else None),
        "count": _to_float(row.get("count")),
        "cost": _to_float(row.get("cost")),
        "product_category": row.get("product_category"),
        "notes": None,
    }
    return out


def _df_to_rows(df: pd.DataFrame, default_currency=None) -> List[dict]:
    df = _map_columns(df)
    rows = []
    for _, r in df.iterrows():
        rows.append(_row_to_schema(r.to_dict(), default_currency=default_currency))
    return rows


# --------- Parsers per filtyp ---------
def _parse_excel(content: bytes) -> List[dict]:
    dfs = pd.read_excel(BytesIO(content), sheet_name=None, engine="openpyxl")
    rows = []
    for _, df in dfs.items():
        rows.extend(_df_to_rows(df))
    return rows


def _parse_csv(content: bytes) -> List[dict]:
    df = pd.read_csv(BytesIO(content), encoding="utf-8", na_filter=False)
    return _df_to_rows(df)


def _parse_docx_to_texts(content: bytes) -> List[str]:
    if docx is None:
        return []
    blobs = []
    doc = docx.Document(BytesIO(content))
    # text från stycken
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if text.strip():
        blobs.append(text)
    # text från tabeller som TSV
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append("\t".join(cell.text.strip() for cell in row.cells))
        tsv = "\n".join(rows).strip()
        if tsv:
            blobs.append(tsv)
    return blobs


def _parse_pdf_to_texts(content: bytes) -> List[str]:
    if pdfplumber is None:
        return []
    blobs = []
    with pdfplumber.open(BytesIO(content)) as pdf:
        pages_text = []
        for page in pdf.pages:
            pages_text.append(page.extract_text() or "")
        text = "\n".join(pages_text).strip()
        if text:
            blobs.append(text)
    return blobs


def parse_attachments(attachments: List[Dict[str, Any]], provider_hint: str = "") -> Dict[str, List]:
    """
    Tar en lista av bilagor (filename, content_type, data) och returnerar:
      {
        "rows":  [ … normaliserade rader (dict) från Excel/CSV … ],
        "texts": [ … textblobs (PDF/DOCX) som kan skickas till LLM … ]
      }
    """
    rows: List[dict] = []
    texts: List[str] = []

    for att in attachments:
        name = (att.get("filename") or "").lower()
        ctype = (att.get("content_type") or "").lower()
        data = att.get("data")

        if not data:
            continue

        try:
            if name.endswith((".xlsx", ".xlsm", ".xls")) or "excel" in ctype:
                rows.extend(_parse_excel(data))
            elif name.endswith(".csv") or "csv" in ctype:
                rows.extend(_parse_csv(data))
            elif name.endswith(".docx") or "word" in ctype:
                texts.extend(_parse_docx_to_texts(data))
            elif name.endswith(".pdf") or "pdf" in ctype:
                texts.extend(_parse_pdf_to_texts(data))
            else:
                # okänd typ – gör inget (eller lägg logik för .zip osv.)
                pass
        except Exception as e:
            print(f"⚠️ Kunde inte tolka bilagan {att.get('filename')}: {e}")

    # sätt provider-hint om möjligt
    for r in rows:
        if not r.get("provider"):
            r["provider"] = provider_hint or None

    return {"rows": rows, "texts": texts}
